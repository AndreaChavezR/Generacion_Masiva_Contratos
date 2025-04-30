import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from data_manager import DataManager
from docx import Document
import os
import sys
from threading import Thread
from datetime import datetime
import pandas as pd

class ContractSystem:
    def __init__(self, root):
        self.root = root
        self.setup_paths()
        self.current_contract_type = tk.StringVar(value='adquisiciones')
        self.data_manager = DataManager(self.current_contract_type.get())
        self.template_path = ""
        self.setup_ui()

    def setup_paths(self):
        """Configura las rutas base y directorios necesarios"""
        if getattr(sys, 'frozen', False):
            self.base_dir = os.path.dirname(sys.executable)
        else:
            self.base_dir = os.getcwd()
        
        self.output_dir = os.path.join(self.base_dir, "contratos_generados")
        self.template_dir = os.path.join(self.base_dir, "plantillas_word")
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)

    def setup_ui(self):
        """Configura la interfaz gráfica principal"""
        self.create_contract_type_selector()
        self.notebook = ttk.Notebook(self.root)
        self.create_tabs()
        self.create_status_bar()

    def create_status_bar(self):
        """Crea la barra de estado en la parte inferior"""
        self.status_bar = ttk.Label(self.root, text="Listo", relief=tk.SUNKEN)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def create_contract_type_selector(self):
        """Crea el selector de tipo de contrato"""
        type_frame = ttk.Frame(self.root)
        type_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(type_frame, text="Tipo de Contrato:").pack(side=tk.LEFT)
        
        self.contract_types = ['adquisiciones', 'servicios']
        type_selector = ttk.Combobox(
            type_frame, 
            textvariable=self.current_contract_type, 
            values=self.contract_types, 
            state='readonly'
        )
        type_selector.pack(side=tk.LEFT, padx=10)
        type_selector.bind('<<ComboboxSelected>>', self.update_contract_type)

    def update_contract_type(self, event=None):
        """Actualiza los componentes al cambiar el tipo de contrato"""
        self.data_manager.switch_contract_type(self.current_contract_type.get())
        self.template_path = ""
        self.clear_form()
        self.rebuild_form()

    def create_tabs(self):
        """Crea las pestañas principales"""
        self.tab_registro = ttk.Frame(self.notebook)
        self.tab_generacion = ttk.Frame(self.notebook)
        
        self.notebook.add(self.tab_registro, text="Registro")
        self.notebook.add(self.tab_generacion, text="Generación")
        self.notebook.pack(expand=True, fill='both')
        
        self.rebuild_form()
        self.create_generation_tab()

    def rebuild_form(self):
        """se reconstruye el formulario según el tipo de contrato"""
        for widget in self.tab_registro.winfo_children():
            widget.destroy()
        self.create_full_registration_tab()

    def create_full_registration_tab(self):
        """Crea los campos del formulario dinámicamente"""
        main_frame = ttk.Frame(self.tab_registro)
        main_frame.pack(fill='both', expand=True)
        
        # Configurar scroll
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        #definir campos según tipo de contrato
        if self.current_contract_type.get() == 'adquisiciones':
        # Lista completa de campos según Excel
            fields = [
                ('NO_CONTRATO', 'entry'),
                ('BIENES', 'text'),
                ('TITULAR_AREA_REQUIRENTE', 'entry'),
                ('TITULAR_AREA', 'entry'),
                ('PROVEEDOR', 'entry'),
                ('NOM_PROVEEDOR', 'entry'),
                ('CARGO_PROVEEDOR', 'entry'),
                ('CARGO_AREA_REQUIRENTE', 'entry'),
                ('FECHA_NOMBRAMIENTO', 'entry'),
                ('FECHA_CELEBRACION', 'text'),
                ('FUNDAMENTO', 'text'),
                ('NO_REQUISICION', 'text'),
                ('TIPO_ADQUISICION', 'entry'),
                ('ADQUISICION', 'text'),
                ('NECESIDADES', 'text'),
                ('PARTIDA_DENOMINACION', 'entry'),
                ('NO_OFICIO', 'entry'),
                ('NO_ESCRITURA_PUBLICA', 'entry'),
                ('FECHA_PUBLICACION', 'entry'),
                ('TITULAR_NOTARIA', 'entry'),
                ('NO_NOTARIA', 'entry'),
                ('NO_MERCANTIL', 'entry'),
                ('ENTIDAD_FEDERATIVA', 'entry'),
                ('DIA', 'entry'),
                ('OBJETO_SOCIAL', 'text'),
                ('PERSONA_FISICA', 'entry'),
                ('CARACTER_PERSONA_FISICA', 'entry'),
                ('IDENTIFICACION', 'entry'),
                ('NO_DOCUMENTO', 'entry'),
                ('INSTITUCION', 'entry'),
                ('FOLIO_REGISTRO_PROVEEDOR', 'entry'),
                ('NO_ESCRITURA', 'entry'),
                ('FECHA_PUBLICACION2', 'entry'),
                ('INE_NOTARIO', 'entry'),
                ('RFC', 'entry'),
                ('NO_CONSTANCIA', 'entry'),
                ('FECHA_EXPEDICION', 'entry'),
                ('ANEXOS', 'text'),
                ('CONSTANCIAS', 'text'),
                ('DOMICILIO', 'entry'),
                ('NUMERO', 'entry'),
                ('COLONIA', 'entry'),
                ('ALCALDIA', 'entry'),
                ('CP', 'entry'),
                ('TELEFONOS', 'entry'),
                ('CORREO', 'entry'),
                ('CALLE', 'entry'),
                ('NO_EXT', 'entry'),
                ('DESCRIPCION_ADQUISICION', 'text'),
                ('NO_REQUERIMIENTO', 'entry'),
                ('PARTIDA_PRESUPUESTAL', 'entry'),
                ('MONTO_AUTORIZADO', 'entry'),
                ('CORREO_1', 'entry'),
                ('CORREO_2', 'entry'),
                ('FECHA_ENTREGA', 'entry'),
                ('FECHA_VIGENCIA_ENTREGA', 'entry'),
                ('VIGENCIA_CONTRATO', 'entry'),
                ('NO_PAG', 'entry'),
                ('DIAS', 'entry'),
                ('MES', 'entry'),
                ('DIA_FIRMA', 'entry'),
                ('MES_FIRMA', 'entry'),
                ('DIRECCION_DE', 'entry')
            ]
        else: 
            fields = [
                ('NO_CONTRATO', 'entry'),
                ('SERVICIOS', 'entry'),
                ('TITULAR_AREA_REQUIRENTE', 'entry'),
                ('TITULAR_AREA', 'entry'),
                ('PROVEEDOR', 'entry'),
                ('NOM_PROVEEDOR', 'entry'),
                ('CARGO_PROVEEDOR', 'entry'),
                ('CARGO_AREA_REQUIRENTE', 'entry'),
                ('FECHA_NOMBRAMIENTO', 'entry'),
                ('TIPO_ADQUISICION', 'entry'),
                ('DESCRIPCION_ADQUISICION', 'entry'),
                ('NO_REQUERIMIENTO', 'entry'),
                ('NECESIDADES', 'entry'),
                ('PARTIDA_DENOMINACION', 'entry'),
                ('NO_OFICIO', 'entry'),
                ('FECHA_NOMBRAMIENTO', 'entry'),
                ('NO_ESCRITURA_PUBLICA', 'entry'),
                ('FECHA_ESCRITURA_PUBLICA', 'entry'),
                ('TITULAR_NOTARIA', 'entry'),
                ('NO_NOTARIA', 'entry'),
                ('NO_MERCANTIL', 'entry'),
                ('ENTIDAD_FEDERATIVA', 'entry'),
                ('INE_NOTARIA', 'entry'),
                ('DIA', 'entry'),
                ('OBJETO_SOCIAL', 'text'),
                ('PERSONA_FISICA', 'entry'),
                ('CARGO_PERSONA_FISICA', 'entry'),
                ('CARGO_REPRESENTANTE', 'entry'),
                ('IDENTIFICACION', 'entry'),
                ('NO_DOCUMENTO', 'entry'),
                ('INSTITUTO', 'entry'),
                ('SEÑALAR_RELACION_CON', 'text'),
                ('NACIONALIDAD', 'entry'),
                ('NO_INE', 'entry'),
                ('EXTRANJERO', 'entry'),
                ('IDENTIFICACION2', 'entry'),
                ('DENOMINACION', 'entry'),
                ('OBJ_SOCIAL', 'text'),
                ('FOLIO', 'entry'),
                ('RFC', 'entry'),
                ('NO_CONSTANCIA', 'entry'),
                ('FECHA_CONSTANCIA', 'entry'),
                ('ANEXOS', 'text'),
                ('CONSTANCIAS', 'text'),
                ('DOMICILIO', 'entry'),
                ('ALCALDIA', 'entry'),
                ('CP', 'entry'),
                ('TELEFONOS', 'entry'),
                ('CORREO', 'entry'),
                ('CALLE', 'entry'),
                ('NO_EXT', 'entry'),
                ('DOMICILIO_CONTRATANTE', 'entry'),
                ('SERVICIO_PROVEEDOR', 'entry'),
                ('NO_SERV', 'entry'),
                ('PARTIDA_PRESUPUESTAL', 'entry'),
                ('MONTO_AUTORIZADO', 'entry'),
                ('CORREO_1', 'entry'),
                ('CORREO_2', 'entry'),
                ('NOM_COORDINADOR', 'entry'),
                ('FECHA_ENTREGA', 'entry'),
                ('FECHA_TERMINO', 'entry'),
                ('FECHA_FIRMA', 'entry'),
                ('NO_PAG', 'entry'),
                ('FECHA_CELEBRACION', 'entry'),
                ('DIRECCION_DE', 'entry')
            ]
   
        self.current_fields = [field[0] for field in fields]

            # Crear campos dinámicamente
        self.entries = {}
        for idx, (field, ftype) in enumerate(fields):
            row = idx % 25
            col = idx // 25
                
            lbl_frame = ttk.Frame(scrollable_frame)
            lbl_frame.grid(row=row, column=col*2, padx=5, pady=2, sticky='w')
                
            lbl = ttk.Label(lbl_frame, text=f"{self.format_label(field)}:")
            lbl.pack(side='left')
                
            entry_frame = ttk.Frame(scrollable_frame)
            entry_frame.grid(row=row, column=col*2+1, padx=5, pady=2, sticky='ew')
                
            if ftype == 'entry':
                entry = ttk.Entry(entry_frame, width=25)
                entry.pack(fill='x')
            else:
                entry = tk.Text(entry_frame, height=3, width=30)
                entry.pack(fill='x')
                
            self.entries[field] = entry
            
        # Configurar scroll y botones
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
            
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
            
        # Botones - FUERA DEL SCROLLABLE FRAME
        btn_frame = ttk.Frame(main_frame)  # Frame principal, no el scrollable
        btn_frame.pack(side=tk.BOTTOM, pady=10)  # Posiciona debajo del scroll
    
        ttk.Button(btn_frame, text="Guardar", command=self.save_data).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Limpiar", command=self.clear_form).pack(side=tk.LEFT, padx=10)

    def create_generation_tab(self):
        """Crea la pestaña de generación de documentos"""
        frame = ttk.Frame(self.tab_generacion, padding=10)
        frame.pack(fill=tk.BOTH, expand=True)

        control_frame = ttk.Frame(frame)
        control_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(control_frame, text="Seleccionar Plantilla", 
                 command=self.select_template).pack(side=tk.LEFT)
        ttk.Button(control_frame, text="Generar Todos", 
                 command=self.start_bulk_generation).pack(side=tk.LEFT, padx=10)
        ttk.Button(control_frame, text="Abrir Carpeta", 
                 command=self.open_output_dir).pack(side=tk.RIGHT)
        
        self.progress = ttk.Progressbar(frame, orient=tk.HORIZONTAL, mode='determinate')
        self.progress.pack(fill=tk.X)

    def format_label(self, text):
        """Formatea los nombres de los campos para mostrar"""
        return text.replace('_', ' ').title()

    def save_data(self):
        data = {}
        for field, entry in self.entries.items():
            if isinstance(entry, tk.Text):
                data[field] = entry.get("1.0", tk.END).strip()
            else:
                data[field] = entry.get().strip()
        
        success, message = self.data_manager.save_record(data)
        if success:
            messagebox.showinfo("Éxito", message)
            self.clear_form()
        else:
            messagebox.showerror("Error", message)

    def clear_form(self):
        for entry in self.entries.values():
            if isinstance(entry, tk.Text):
                entry.delete("1.0", tk.END)
            else:
                entry.delete(0, tk.END)

    def select_template(self):
        contract_type = self.current_contract_type.get()
        initial_dir = os.path.join(self.template_dir, contract_type)
        os.makedirs(initial_dir, exist_ok=True)
        
        # Solo permitir .docx
        self.template_path = filedialog.askopenfilename(
            initialdir=initial_dir,
            filetypes=[("Plantillas Word", "*.docx *doc")],
            title=f"Seleccionar plantilla de {contract_type}"
        )
        if self.template_path:
            self.update_status(f"Plantilla: {os.path.basename(self.template_path)}")

    
    def start_bulk_generation(self):
        if not self.template_path:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla primero")
            return
        
        Thread(target=self.generate_all_documents, daemon=True).start()

    def generate_all_documents(self):
        """Genera los documentos según el tipo de contrato"""
        try:
            # Validación crítica antes de comenzar
            if not self.template_path:
                raise ValueError("No se ha seleccionado ninguna plantilla")
                
            # Normalizar rutas y verificar existencia
            template_path = os.path.normpath(self.template_path)
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Archivo de plantilla no encontrado: {template_path}")

            # Obtener registros pendientes
            df = self.data_manager.get_pending_records()
            total = len(df)
            
            if total == 0:
                messagebox.showinfo("Información", "No hay contratos pendientes")
                return

            # Configurar progreso
            self.progress['maximum'] = total
            os.makedirs(self.output_dir, exist_ok=True)
            
            # Obtener extensión del archivo plantilla
            file_extension = os.path.splitext(template_path)[1]

            for idx, (_, row) in enumerate(df.iterrows()):
                # Cargar plantilla
                doc = Document(template_path)
                replacements = row.to_dict()
            
                # Agregar asterisco a los campos capturados por el usuario
                for field in self.current_fields:  # Usar la lista de campos actual
                    if field in replacements:
                        replacements[field] = f"*{replacements[field]}*"    

                # Campos adicionales dinámicos
                replacements['FECHA_GENERACION'] = datetime.now().strftime("%d/%m/%Y %H:%M")

                # Reemplazo en todo el documento
                self.replace_template_content(doc, replacements)

                # Generar nombre de archivo seguro
                contract_number = str(row['NO_CONTRATO']).replace('/', '_').strip()
                filename = f"Contrato_{contract_number}{file_extension}"
                output_path = os.path.join(self.output_dir, filename)
                
                # Guardar documento
                doc.save(output_path)
                
                # Actualizar estado
                self.data_manager.mark_as_generated(row['ID'])
                self.progress['value'] = idx + 1
                self.update_status(f"Generado: {filename}")
            
            messagebox.showinfo("Éxito", f"Se generaron {total} contratos exitosamente")
            self.progress['value'] = 0
            
        except Exception as e:
            messagebox.showerror("Error Crítico", f"Fallo en generación: {str(e)}")
            self.progress['value'] = 0

    def replace_template_content(self, doc, replacements):
        """Realiza el reemplazo de variables en toda la plantilla"""
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if f'{{{{{key}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Reemplazar en encabezados y pies de página
        for section in doc.sections:
            self.replace_in_header_footer(section.header, replacements)
            self.replace_in_header_footer(section.footer, replacements)

    def replace_in_header_footer(self, header_footer, replacements):
        """Reemplaza contenido en encabezados y pies de página"""
        for paragraph in header_footer.paragraphs:
            for key, value in replacements.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

    def open_output_dir(self):
        try:
            os.startfile(self.output_dir)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta: {str(e)}")

    def update_status(self, message):
        self.status_bar.config(text=message)
        self.root.update_idletasks()

if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("1400x900")
    app = ContractSystem(root)
    root.mainloop()
